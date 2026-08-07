"""Create a small sample DOCX for local demos."""

from pathlib import Path

from docx import Document


def main() -> None:
    out = Path("sample_docs/sample_service_agreement.docx")
    out.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    doc.add_heading("Service Agreement", level=1)
    doc.add_paragraph(
        "This Service Agreement (the \"Agreement\") is entered into between Acme Corp "
        "(\"Provider\") and the Client. The parties agree to the following clauses."
    )
    doc.add_heading("1. Term and Termination", level=2)
    doc.add_paragraph(
        "Either party may terminate this Agreement by providing thirty (30) days written "
        "notice. Provider may terminate immediately for material breach that remains "
        "uncured for fifteen (15) days after notice. Upon termination, Client shall pay "
        "all outstanding invoices within fourteen (14) days."
    )
    doc.add_heading("2. Limitation of Liability", level=2)
    doc.add_paragraph(
        "Except for willful misconduct, Provider's aggregate liability under this Agreement "
        "shall not exceed the fees paid by Client in the twelve (12) months preceding the claim. "
        "Neither party shall be liable for indirect or consequential damages."
    )
    doc.add_heading("3. Governing Law", level=2)
    doc.add_paragraph(
        "This Agreement shall be governed by the laws of the State of Delaware, without "
        "regard to conflict of law principles. Disputes shall be resolved in the courts "
        "located in Wilmington, Delaware."
    )
    doc.add_heading("4. Indemnification", level=2)
    doc.add_paragraph(
        "Client shall indemnify and hold harmless Provider from claims arising out of "
        "Client's misuse of the services, except to the extent caused by Provider's negligence."
    )
    doc.save(out)
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
